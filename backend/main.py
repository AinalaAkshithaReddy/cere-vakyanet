from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from deep_translator import GoogleTranslator
import pdfplumber
import langdetect
from typing import Optional, List, Dict
from pathlib import Path
import io
import os
import logging
import traceback
import re
import time
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from reportlab.lib.colors import HexColor, black
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping
from font_utils import (
    UNICODE_FONT_NAME, UNICODE_FONT_BOLD, get_font_for_text, 
    detect_telugu_text, detect_script, get_script_from_language
)

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI(title="Translation API")

# CORS middleware to allow frontend to connect
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000","https://cere-vakyanet.cerevyn.com"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
# USE_MOCK_TRANSLATION = os.getenv("USE_MOCK_TRANSLATION", "false").lower() == "true"
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
# AZURE_TRANSLATOR_KEY = os.getenv("AZURE_TRANSLATOR_KEY", "")
# AZURE_TRANSLATOR_REGION = os.getenv("AZURE_TRANSLATOR_REGION", "")

# Log configuration status
# logger.info(f"Mock translation mode: {USE_MOCK_TRANSLATION}")
# if USE_MOCK_TRANSLATION:
#     logger.info("⚠️  Running in MOCK mode - translations will be simulated")
if GOOGLE_API_KEY:
    logger.info("✓ Google API key found")
else: 
    logger.warning("⚠️  Google API key not found")
# if AZURE_TRANSLATOR_KEY:
#     logger.info("✓ Azure Translator key found")

# Language mapping
LANGUAGE_MAP = {
    "hindi": "hi",
    "english": "en",
    "urdu": "ur",
    "assamese": "as",
    "bengali": "bn",
    "gujarati": "gu",
    "kannada": "kn",
    "malayalam": "ml",
    "marathi": "mr",
    "odia": "or",
    "punjabi": "pa",
    "tamil": "ta",
    "telugu": "te"
}

# Reverse mapping for display names
LANGUAGE_DISPLAY = {
    "hi": "Hindi",
    "en": "English",
    "ur": "Urdu",
    "as": "Assamese",
    "bn": "Bengali",
    "gu": "Gujarati",
    "kn": "Kannada",
    "ml": "Malayalam",
    "mr": "Marathi",
    "or": "Odia",
    "pa": "Punjabi",
    "ta": "Tamil",
    "te": "Telugu"
}

# Mock translation responses for testing
# MOCK_TRANSLATIONS = {
#     "hi": {"hello": "नमस्ते", "thank you": "धन्यवाद", "good morning": "सुप्रभात"},
#     "en": {"hello": "hello", "thank you": "thank you", "good morning": "good morning"},
#     "ta": {"hello": "வணக்கம்", "thank you": "நன்றி", "good morning": "காலை வணக்கம்"},
# }

class TranslateRequest(BaseModel):
    text: str
    target_language: str

class ParagraphModel(BaseModel):
    id: str
    text: str

class ParagraphTranslateRequest(BaseModel):
    paragraphs: List[ParagraphModel]
    targetLang: str

class TranslationResult(BaseModel):
    id: str
    translatedText: str
    confidence: Optional[str] = None

class ParagraphTranslateResponse(BaseModel):
    translations: List[TranslationResult]
    provider: str
    error: Optional[str] = None

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF file"""
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception as e:
        error_msg = f"Error extracting text from PDF: {str(e)}"
        logger.error(f"PDF extraction error: {traceback.format_exc()}")
        raise HTTPException(status_code=400, detail=error_msg)

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX file (including headings, paragraphs, and tables)"""
    try:
        # Validate DOCX file
        if len(docx_bytes) < 100:
            raise ValueError("DOCX file appears to be too small or corrupted")
        
        # Create BytesIO buffer and extract text
        buffer = io.BytesIO(docx_bytes)
        doc = Document(buffer)
        
        text_parts = []
        
        # Extract text from all paragraphs (including headings)
        # Skip the title if it matches "Translated Text (Language)" pattern
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                # Skip the title heading if it matches the pattern we add during generation
                # Check if this is a heading that matches our title pattern
                if paragraph.style.name.startswith('Heading'):
                    if para_text.startswith('Translated Text (') and para_text.endswith(')'):
                        continue  # Skip the auto-generated title
                text_parts.append(para_text)
        
        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        
        # Join all text parts
        extracted_text = "\n".join(text_parts).strip()
        
        if not extracted_text:
            logger.warning("No text extracted from DOCX file - file may be empty or contain only images")
            raise ValueError("No text content found in DOCX file. The file may contain only images or be empty.")
        
        # Clean up the extracted text
        # Remove excessive whitespace and normalize line breaks
        # Replace multiple newlines with double newline (paragraph break)
        extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text)
        # Replace multiple spaces with single space (but preserve intentional spacing)
        extracted_text = re.sub(r'[ \t]+', ' ', extracted_text)
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in extracted_text.split('\n')]
        extracted_text = '\n'.join(line for line in lines if line)  # Remove empty lines
        
        if not extracted_text:
            logger.warning("No text remaining after cleaning DOCX file")
            raise ValueError("No text content found in DOCX file after processing.")
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
        
    except ValueError as e:
        # Re-raise ValueError as-is (these are expected errors)
        error_msg = f"Error extracting text from DOCX: {str(e)}"
        logger.error(f"DOCX extraction error: {error_msg}")
        raise HTTPException(status_code=400, detail=error_msg)
    except Exception as e:
        error_msg = f"Error extracting text from DOCX: {str(e)}"
        logger.error(f"DOCX extraction error: {traceback.format_exc()}")
        raise HTTPException(status_code=400, detail=error_msg)

def extract_text_from_txt(txt_bytes: bytes) -> str:
    """Extract text from TXT file"""
    try:
        # Try to decode with common encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text = None
        for encoding in encodings:
            try:
                text = txt_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            # Fallback to utf-8 with errors='ignore'
            text = txt_bytes.decode('utf-8', errors='ignore')
        
        return text.strip()
    except Exception as e:
        error_msg = f"Error extracting text from TXT: {str(e)}"
        logger.error(f"TXT extraction error: {traceback.format_exc()}")
        raise HTTPException(status_code=400, detail=error_msg)

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from file based on extension"""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith('.txt'):
        return extract_text_from_txt(file_bytes)
    else:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Please upload a PDF, DOCX, or TXT file."
        )

def detect_language(text: str) -> str:
    """Detect the language of the input text"""
    try:
        if not text or not text.strip():
            return "Unknown"
        detected = langdetect.detect(text)
        return LANGUAGE_DISPLAY.get(detected, detected.upper())
    except Exception as e:
        logger.warning(f"Language detection failed: {str(e)}")
        return "Unknown"

def mock_translate(text: str, target_language: str) -> str:
    """Mock translation for testing"""
    target_code = LANGUAGE_MAP.get(target_language.lower(), "hi")
    text_lower = text.lower().strip()
    
    # Try to find exact match in mock translations
    # if text_lower in MOCK_TRANSLATIONS.get(target_code, {}):
    #     return MOCK_TRANSLATIONS[target_code][text_lower]
    
    # Return a mock translation with target language indicator
    return f"[MOCK: {target_code.upper()}] {text}"

def translate_text(text: str, target_language: str) -> Dict[str, any]:
    """Translate text to target language with detailed error handling"""
    try:
        # Validate input
        if not text or not text.strip():
            raise ValueError("Text to translate cannot be empty")
        
        target_code = LANGUAGE_MAP.get(target_language.lower())
        if not target_code:
            raise ValueError(f"Unsupported target language: {target_language}")
        
        # Mock mode
        # if USE_MOCK_TRANSLATION:
        #     logger.info(f"MOCK: Translating '{text[:50]}...' to {target_code}")
        #     translated = mock_translate(text, target_language)
        #     return {
        #         "translated_text": translated,
        #         "confidence": "high",
        #         "provider": "mock"
        #     }
        
        # Detect source language with better error handling
        source_code = "auto"  # Default to auto
        try:
            detected = langdetect.detect(text)
            # Validate detected language code
            if detected and len(detected) == 2:
                source_code = detected
                logger.info(f"Detected source language: {source_code}")
            else:
                logger.warning(f"Invalid language detection result: {detected}, using 'auto'")
                source_code = "auto"
        except Exception as e:
            logger.warning(f"Language detection failed, using 'auto': {str(e)}")
            source_code = "auto"
        
        # If source and target are same, return original
        if source_code == target_code:
            logger.info("Source and target languages are the same, returning original text")
            return {
                "translated_text": text,
                "confidence": "high",
                "provider": "google"
            }
        
        # Attempt translation with Google Translator (with retry logic)
        max_retries = 2
        last_error = None
        
        for attempt in range(max_retries):
            try:
                logger.info(f"Translating from {source_code} to {target_code} (attempt {attempt + 1}/{max_retries})")
                
                # Always use 'auto' for source to let Google handle detection more reliably
                # This is safer for Indian languages
                translator = GoogleTranslator(source='auto', target=target_code)
                translated = translator.translate(text)
                
                if not translated or not translated.strip():
                    raise ValueError("Translation returned empty result")
                
                logger.info(f"Translation successful: '{text[:30]}...' -> '{translated[:30]}...'")
                return {
                    "translated_text": translated,
                    "confidence": "high",
                    "provider": "google"
                }
                
            except Exception as translate_error:
                last_error = translate_error
                error_msg = str(translate_error)
                error_type = type(translate_error).__name__
                
                logger.warning(f"Translation attempt {attempt + 1} failed ({error_type}): {error_msg}")
                
                # If it's a connection/timeout error and we have retries left, wait and retry
                if attempt < max_retries - 1:
                    if "timeout" in error_msg.lower() or "connection" in error_msg.lower() or "503" in error_msg or "unavailable" in error_msg.lower():
                        wait_time = (attempt + 1) * 1  # Wait 1s, then 2s
                        logger.info(f"Retrying after {wait_time} seconds...")
                        time.sleep(wait_time)
                        continue
                
                # If we're here, either no retries left or it's a different error
                break
        
        # Handle the final error
        if last_error:
            translate_error = last_error
            error_msg = str(translate_error)
            error_type = type(translate_error).__name__
            
            # Log detailed error information
            logger.error(f"Translation error ({error_type}): {error_msg}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            
            # Provide user-friendly error messages
            error_msg_lower = error_msg.lower()
            
            if "429" in error_msg or "rate limit" in error_msg_lower or "quota" in error_msg_lower:
                raise HTTPException(
                    status_code=429,
                    detail="Translation rate limit exceeded. Please wait a moment and try again."
                )
            elif any(keyword in error_msg_lower for keyword in ["timeout", "connection", "503", "service unavailable", "temporarily unavailable", "network", "unreachable"]):
                raise HTTPException(
                    status_code=503,
                    detail=f"Translation service is temporarily unavailable. Error: {error_msg}. Please check your internet connection and try again in a moment."
                )
            elif any(keyword in error_msg_lower for keyword in ["invalid", "not supported", "unsupported", "language"]):
                raise HTTPException(
                    status_code=400,
                    detail=f"Translation failed: {error_msg}. Please check the language selection and try again."
                )
            elif "empty" in error_msg_lower or "no result" in error_msg_lower:
                raise HTTPException(
                    status_code=400,
                    detail="Translation returned empty result. Please check the input text and try again."
                )
            else:
                # Return detailed error for debugging
                logger.error(f"Unexpected translation error: {error_type} - {error_msg}")
                raise HTTPException(
                    status_code=502,
                    detail=f"Translation service error: {error_msg[:200]}. Please try again. If this persists, the translation service may be experiencing issues."
                )
                
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        error_msg = f"Unexpected error during translation: {str(e)}"
        logger.error(f"Unexpected error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=error_msg)

@app.get("/")
def read_root():
    """Health check endpoint"""
    return {
        "message": "Translation API is running",
        # "mock_mode": USE_MOCK_TRANSLATION,
        # "provider": "mock" if USE_MOCK_TRANSLATION else "google"
    }

@app.get("/api/health")
def health_check():
    """Detailed health check"""
    return {
        "status": "healthy",
        "has_google_key": bool(GOOGLE_API_KEY),
    }

@app.post("/api/detect-language")
async def detect_language_endpoint(request: TranslateRequest):
    """Detect language of input text"""
    try:
        detected = detect_language(request.text)
        return {"detected_language": detected}
    except Exception as e:
        logger.error(f"Language detection error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Language detection failed: {str(e)}")

@app.post("/api/translate")
async def translate_endpoint(request: TranslateRequest):
    """Translate text to target language"""
    try:
        logger.info(f"Translation request: '{request.text[:50]}...' -> {request.target_language}")
        detected_lang = detect_language(request.text)
        result = translate_text(request.text, request.target_language)
        
        return {
            "original_text": request.text,
            "translated_text": result["translated_text"],
            "detected_language": detected_lang,
            "target_language": request.target_language,
            "confidence": result.get("confidence", "high"),
            "provider": result.get("provider", "google")
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Translation endpoint error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/api/translate-paragraphs")
async def translate_paragraphs_endpoint(request: ParagraphTranslateRequest):
    """Translate multiple paragraphs with IDs"""
    try:
        logger.info(f"Paragraph translation request: {len(request.paragraphs)} paragraphs -> {request.targetLang}")
        translations = []
        errors = []
        
        for para in request.paragraphs:
            try:
                result = translate_text(para.text, request.targetLang)
                translations.append(TranslationResult(
                    id=para.id,
                    translatedText=result["translated_text"],
                    confidence=result.get("confidence", "medium")
                ))
            except Exception as e:
                error_msg = str(e)
                errors.append(f"Paragraph {para.id}: {error_msg}")
                logger.error(f"Error translating paragraph {para.id}: {error_msg}")
                # Still add a result but with error indication
                translations.append(TranslationResult(
                    id=para.id,
                    translatedText=f"[Error: {error_msg}]",
                    confidence="low"
                ))
        
        return ParagraphTranslateResponse(
            translations=translations,
            # provider="mock" if USE_MOCK_TRANSLATION else "google",
            error="; ".join(errors) if errors else None
        )
    except Exception as e:
        logger.error(f"Paragraph translation endpoint error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Paragraph translation failed: {str(e)}")

@app.post("/api/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    """Upload and extract text from PDF, DOCX, or TXT file"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        filename_lower = file.filename.lower()
        supported_extensions = ['.pdf', '.docx', '.txt']
        
        if not any(filename_lower.endswith(ext) for ext in supported_extensions):
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Please upload a PDF, DOCX, or TXT file."
            )
        
        file_type = "PDF" if filename_lower.endswith('.pdf') else "DOCX" if filename_lower.endswith('.docx') else "TXT"
        logger.info(f"{file_type} upload: {file.filename}")
        contents = await file.read()
        extracted_text = extract_text_from_file(contents, file.filename)
        
        if not extracted_text:
            raise HTTPException(
                status_code=400, 
                detail=f"No text could be extracted from {file_type} file"
            )
        
        detected_lang = detect_language(extracted_text)
        logger.info(f"{file_type} text extracted: {len(extracted_text)} characters, detected language: {detected_lang}")
        
        return {
            "extracted_text": extracted_text,
            "detected_language": detected_lang,
            "file_type": file_type
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File upload error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"File processing failed: {str(e)}")

@app.post("/api/translate-pdf")
async def translate_pdf(file: UploadFile = File(...), target_language: str = Query(default="hindi")):
    """Upload PDF, DOCX, or TXT file, extract text, and translate"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        filename_lower = file.filename.lower()
        supported_extensions = ['.pdf', '.docx', '.txt']
        
        if not any(filename_lower.endswith(ext) for ext in supported_extensions):
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Please upload a PDF, DOCX, or TXT file."
            )
        
        file_type = "PDF" if filename_lower.endswith('.pdf') else "DOCX" if filename_lower.endswith('.docx') else "TXT"
        logger.info(f"{file_type} translate: {file.filename} -> {target_language}")
        
        contents = await file.read()
        logger.info(f"File size: {len(contents)} bytes")
        
        extracted_text = extract_text_from_file(contents, file.filename)
        logger.info(f"Extracted text length: {len(extracted_text)} characters")
        
        if not extracted_text:
            raise HTTPException(
                status_code=400, 
                detail=f"No text could be extracted from {file_type} file. The file may be empty or contain only images."
            )
        
        detected_lang = detect_language(extracted_text)
        logger.info(f"Detected language: {detected_lang}")
        
        result = translate_text(extracted_text, target_language)
        logger.info(f"Translation successful: {len(result['translated_text'])} characters")
        
        return {
            "original_text": extracted_text,
            "translated_text": result["translated_text"],
            "detected_language": detected_lang,
            "target_language": target_language,
            "confidence": result.get("confidence", "medium"),
            "provider": result.get("provider", "google"),
            "file_type": file_type
        }
    except HTTPException:
        raise
    except Exception as e:
        error_detail = str(e)
        logger.error(f"File translation error: {error_detail}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=500, 
            detail=f"File translation failed: {error_detail}. Please check the file format and try again."
        )

def generate_docx_from_text(text: str, title: str = "Translated Text", target_language: str = "") -> bytes:
    """
    Generate a DOCX file from translated text with proper Unicode font support for Indian scripts.
    Uses appropriate fonts for Telugu, Malayalam, Tamil, and Devanagari scripts.
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Set default font for the document (will be overridden per paragraph if needed)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # Default font
        font.size = Pt(12)
        
        # Get script information
        script = get_script_from_language(target_language) or detect_script(text)
        is_indian_script = script in ['telugu', 'malayalam', 'tamil', 'devanagari']
        
        # Font mapping for Indian scripts (using system fonts that support Unicode)
        font_mapping = {
            'telugu': 'Noto Sans Telugu',  # Will fallback to system fonts if not available
            'malayalam': 'Noto Sans Malayalam',
            'tamil': 'Noto Sans Tamil',
            'devanagari': 'Noto Sans Devanagari',
        }
        
        # Determine font size
        if is_indian_script:
            body_font_size = 15
            title_font_size = 18
            font_family = font_mapping.get(script, 'Arial Unicode MS')
        else:
            body_font_size = 12
            title_font_size = 16
            font_family = 'Calibri'
        
        # Add title
        if target_language:
            title_text = f"Translated Text ({target_language.title()})"
        else:
            title_text = title
        
        title_para = doc.add_heading(title_text, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.runs[0]
        title_run.font.size = Pt(title_font_size)
        title_run.font.bold = True
        
        # Set fonts for complex scripts
        if is_indian_script:
            # Set both Latin and East Asian fonts for proper rendering
            rFonts = title_run._element.rPr.rFonts
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                title_run._element.rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_family)
            rFonts.set(qn('w:hAnsi'), font_family)
            rFonts.set(qn('w:eastAsia'), font_family)
            rFonts.set(qn('w:cs'), font_family)
        
        # Add spacing after title
        doc.add_paragraph()  # Empty paragraph for spacing
        
        # Add content paragraphs
        paragraphs = text.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:  # Skip empty lines
                para = doc.add_paragraph(para_text)
                para_format = para.paragraph_format
                para_format.space_after = Pt(12)  # Space after paragraph
                para_format.line_spacing = 1.5  # 1.5x line spacing
                
                # Apply formatting to all runs in the paragraph
                for run in para.runs:
                    run.font.size = Pt(body_font_size)
                    run.font.bold = False  # Regular weight for body text
                    
                    # Set fonts for complex scripts
                    if is_indian_script:
                        # Set both Latin and East Asian fonts for proper rendering
                        rFonts = run._element.rPr.rFonts
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            run._element.rPr.append(rFonts)
                        rFonts.set(qn('w:ascii'), font_family)
                        rFonts.set(qn('w:hAnsi'), font_family)
                        rFonts.set(qn('w:eastAsia'), font_family)
                        rFonts.set(qn('w:cs'), font_family)
        
        # Save document to BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"DOCX generation error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

def generate_pdf_from_text(text: str, title: str = "Translated Text", target_language: str = "") -> bytes:
    """
    Generate a PDF file from translated text with proper Unicode font support for Indian scripts.
    Uses appropriate fonts (Noto Sans) for Telugu, Malayalam, Tamil, and Devanagari scripts.
    """
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
        
        # Container for the 'Flowable' objects
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Get appropriate font for the target language/text
        font_name, use_bold = get_font_for_text(text, target_language)
        
        # Determine if this is an Indian script that needs enhanced styling
        script = get_script_from_language(target_language) or detect_script(text)
        is_indian_script = script in ['telugu', 'malayalam', 'tamil', 'devanagari']
        
        # Enhanced styling for Indian scripts: larger font, better spacing
        if is_indian_script:
            # Use larger font size (14-16pt) for better readability
            body_font_size = 15
            # Line height = 1.5x font size for comfortable reading
            body_leading = int(body_font_size * 1.5)  # ~22.5, rounded to 23
            body_font_name = font_name
            logger.info(f"Using {script} font: {body_font_name} at size {body_font_size}pt with {body_leading}pt leading")
        else:
            # Standard styling for other languages
            body_font_size = 12
            body_leading = int(body_font_size * 1.5)  # 18pt
            body_font_name = font_name if font_name != UNICODE_FONT_NAME else UNICODE_FONT_NAME
        
        # Create title style with proper Unicode font
        title_style = ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=HexColor('#1f2937'),
            spaceAfter=24,
            alignment=TA_LEFT,
            fontName=body_font_name if is_indian_script else UNICODE_FONT_BOLD
        )
        
        # Add title
        if target_language:
            title_text = f"Translated Text ({target_language.title()})"
        else:
            title_text = title
        # Escape HTML special characters for title
        title_escaped = title_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(title_escaped, title_style))
        story.append(Spacer(1, 12))
        
        # Create body style with proper formatting
        body_style = ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=body_font_size,
            textColor=HexColor('#111827'),
            leading=body_leading,
            alignment=TA_LEFT,
            fontName=body_font_name,
            spaceAfter=12,  # Space between paragraphs
            leftIndent=0,
            rightIndent=0,
        )
        
        # Add content with proper paragraph handling
        # Split text into paragraphs (by newlines) and process each
        paragraphs = text.split('\n')
        for para in paragraphs:
            para = para.strip()
            if para:  # Skip empty lines
                # Escape HTML special characters
                para_text = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                # Replace multiple spaces with single space (but preserve intentional formatting)
                para_text = ' '.join(para_text.split())
                story.append(Paragraph(para_text, body_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"PDF generation error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

class DownloadDocxRequest(BaseModel):
    text: str
    title: Optional[str] = "Translated Text"
    target_language: Optional[str] = ""
    original_filename: Optional[str] = ""  # Original filename for DOCX naming

@app.post("/api/download-docx")
async def download_docx(request: DownloadDocxRequest):
    """Generate and download DOCX from translated text"""
    try:
        if not request.text or not request.text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        logger.info(f"Generating DOCX for {len(request.text)} characters, target language: {request.target_language}")
        docx_bytes = generate_docx_from_text(
            request.text, 
            request.title or "Translated Text",
            request.target_language or ""
        )
        
        # Create filename: originalname_targetLang_YYYYMMDD.docx
        from datetime import datetime
        
        # Get base name from original filename (without extension)
        if request.original_filename:
            base_name = Path(request.original_filename).stem
            # Remove any special characters that might cause issues
            base_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_')).strip()
            if not base_name:
                base_name = "translated_text"
        else:
            base_name = "translated_text"
        
        # Get target language code (lowercase, no spaces)
        target_lang = ""
        if request.target_language:
            target_lang = request.target_language.lower().replace(' ', '_')
        
        # Get current date in YYYYMMDD format
        date_str = datetime.now().strftime("%Y%m%d")
        
        # Build filename
        filename_parts = [base_name]
        if target_lang:
            filename_parts.append(target_lang)
        filename_parts.append(date_str)
        filename = "_".join(filename_parts) + ".docx"
        
        logger.info(f"Generated DOCX filename: {filename}")
        
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX download error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"DOCX download failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
